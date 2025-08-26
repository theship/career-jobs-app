"""Resume processing service with multi-stage skill extraction."""

import io
import logging
import re
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

import numpy as np
import openai
from docx import Document
from pdfminer.high_level import extract_text
from pdfminer.layout import LAParams
from rapidfuzz import fuzz, process

from ..models.resumes import SkillExtractionResult
from ..utils.config import settings

logger = logging.getLogger(__name__)


class ResumeProcessor:
    """Process resumes with multi-stage skill extraction pipeline."""

    def __init__(self):
        """Initialize the resume processor."""
        self.skills_vocab = self._load_skills_vocabulary()
        self.openai_client = openai.AsyncClient(api_key=settings.openai_api_key)

    def _load_skills_vocabulary(self) -> Dict[str, Dict[str, Any]]:
        """Load skills vocabulary from CSV."""
        import csv
        import os

        vocab = {}
        csv_path = os.path.join(settings.CONFIG_DIR, "skills_vocab.csv")

        if not os.path.exists(csv_path):
            logger.warning(
                f"Skills vocabulary not found at {csv_path}, using empty vocabulary"
            )
            return vocab

        try:
            with open(csv_path, "r", encoding="utf-8") as f:
                reader = csv.DictReader(f)
                for row in reader:
                    skill = row.get("skill", "").strip()
                    if skill:
                        vocab[skill.lower()] = {
                            "canonical": skill,
                            "category": row.get("category", "").strip() or "Other",
                            "aliases": [
                                a.strip()
                                for a in (row.get("aliases", "") or "").split("|")
                                if a.strip()
                            ],
                            "tags": [
                                t.strip()
                                for t in (row.get("tags", "") or "").split("|")
                                if t.strip()
                            ],
                        }
                        # Add aliases as keys pointing to the same skill
                        for alias in vocab[skill.lower()]["aliases"]:
                            if alias.lower() not in vocab:
                                vocab[alias.lower()] = vocab[skill.lower()]
        except Exception as e:
            logger.error(f"Failed to load skills vocabulary: {e}")

        logger.info(f"Loaded {len(vocab)} skills from vocabulary")
        return vocab

    async def extract_text(self, file: Any) -> str:
        """Extract text from uploaded file (PDF, DOCX, or TXT)."""
        filename = file.filename.lower()
        content = await file.read()

        try:
            if filename.endswith(".pdf"):
                return self._extract_pdf_text(content)
            elif filename.endswith(".docx"):
                return self._extract_docx_text(content)
            elif filename.endswith(".txt"):
                return content.decode("utf-8", errors="ignore")
            else:
                raise ValueError(f"Unsupported file type: {filename}")
        finally:
            # Reset file pointer for potential reuse
            await file.seek(0)

    def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF content."""
        try:
            laparams = LAParams(
                line_overlap=0.5,
                char_margin=2.0,
                word_margin=0.1,
                boxes_flow=0.5,
                detect_vertical=False,
            )
            text = extract_text(
                io.BytesIO(content),
                laparams=laparams,
                maxpages=20,  # Limit to 20 pages for performance
            )
            return text
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise ValueError(f"Failed to extract text from PDF: {e}")

    def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX content."""
        try:
            doc = Document(io.BytesIO(content))
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)

            return "\n".join(paragraphs)
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise ValueError(f"Failed to extract text from DOCX: {e}")

    async def extract_skills(self, text: str) -> SkillExtractionResult:
        """
        Multi-stage skill extraction pipeline.

        Stages:
        1. Dictionary/fuzzy matching with rapidfuzz
        2. Embedding-based candidate retrieval (if coverage < 70%)
        3. OpenAI function calling with closed-world constraint (if still needed)
        4. Pydantic validation and merging
        """
        all_skills = {}
        evidence_spans = {}
        confidence_scores = {}
        years_experience = {}

        # Stage 1: Dictionary/fuzzy matching
        stage1_skills = self._extract_skills_fuzzy(text)
        for skill, data in stage1_skills.items():
            all_skills[skill] = data["canonical"]
            evidence_spans[skill] = data["spans"]
            confidence_scores[skill] = data["confidence"]
            if "years" in data:
                years_experience[skill] = data["years"]

        # Calculate coverage
        coverage = len(all_skills) / max(len(self.skills_vocab), 1) * 100
        method = "fuzzy_matching"

        # Stage 2: Embedding-based retrieval if coverage < 70%
        if coverage < 70 and settings.openai_api_key:
            try:
                stage2_skills = await self._extract_skills_embeddings(text, all_skills)
                for skill, data in stage2_skills.items():
                    if skill not in all_skills:
                        all_skills[skill] = data["canonical"]
                        evidence_spans[skill] = data["spans"]
                        confidence_scores[skill] = data["confidence"]
                        if "years" in data:
                            years_experience[skill] = data["years"]

                coverage = len(all_skills) / max(len(self.skills_vocab), 1) * 100
                method = "fuzzy_and_embeddings"
            except Exception as e:
                logger.warning(f"Embedding extraction failed: {e}")

        # Stage 3: OpenAI function calling if still low coverage
        if coverage < 70 and settings.openai_api_key:
            try:
                stage3_skills = await self._extract_skills_openai(text, all_skills)
                for skill, data in stage3_skills.items():
                    if skill not in all_skills:
                        all_skills[skill] = data["canonical"]
                        evidence_spans[skill] = data.get("spans", [])
                        confidence_scores[skill] = data.get("confidence", 0.7)
                        if "years" in data:
                            years_experience[skill] = data["years"]

                coverage = len(all_skills) / max(len(self.skills_vocab), 1) * 100
                method = "full_pipeline"
            except Exception as e:
                logger.warning(f"OpenAI extraction failed: {e}")

        # Filter out None values from years_experience
        filtered_years = (
            {k: v for k, v in years_experience.items() if v is not None}
            if years_experience
            else None
        )

        return SkillExtractionResult(
            skills=list(all_skills.values()),
            method=method,
            confidence_scores=confidence_scores,
            evidence_spans=evidence_spans,
            coverage=min(coverage, 100.0),
            years_experience=filtered_years if filtered_years else None,
        )

    def _extract_skills_fuzzy(self, text: str) -> Dict[str, Dict[str, Any]]:
        """Stage 1: Extract skills using fuzzy matching."""
        found_skills = {}
        text_lower = text.lower()

        # Split text into words for context
        words = re.findall(r"\b\w+\b", text_lower)
        word_positions = {}
        current_pos = 0

        for word in words:
            pos = text_lower.find(word, current_pos)
            if pos != -1:
                word_positions[word] = pos
                current_pos = pos + len(word)

        # Search for each skill in vocabulary
        for skill_key, skill_data in self.skills_vocab.items():
            canonical = skill_data["canonical"]

            # Direct match
            if skill_key in text_lower:
                spans = self._find_spans(text, skill_key)
                if spans:
                    years = self._extract_years_for_skill(text, canonical)
                    skill_data = {
                        "canonical": canonical,
                        "spans": spans,
                        "confidence": 1.0,
                    }
                    if years is not None:
                        skill_data["years"] = years
                    found_skills[canonical] = skill_data
                    continue

            # Fuzzy match for multi-word skills
            if " " in skill_key:
                # Use fuzzy matching for phrases
                ratio = fuzz.partial_ratio(skill_key, text_lower)
                if ratio > 85:  # High threshold for multi-word matches
                    # Find approximate location
                    spans = self._find_fuzzy_spans(text, skill_key)
                    if spans:
                        years = self._extract_years_for_skill(text, canonical)
                        skill_data = {
                            "canonical": canonical,
                            "spans": spans,
                            "confidence": ratio / 100.0,
                        }
                        if years is not None:
                            skill_data["years"] = years
                        found_skills[canonical] = skill_data

        return found_skills

    def _find_spans(self, text: str, skill: str) -> List[Dict[str, int]]:
        """Find exact character spans for a skill in text."""
        spans = []
        pattern = re.compile(r"\b" + re.escape(skill) + r"\b", re.IGNORECASE)

        for match in pattern.finditer(text):
            spans.append({"start": match.start(), "end": match.end()})

        return spans[:3]  # Return up to 3 evidence spans

    def _find_fuzzy_spans(self, text: str, skill: str) -> List[Dict[str, int]]:
        """Find approximate character spans for a skill using fuzzy matching."""
        spans = []
        text_lower = text.lower()
        skill_lower = skill.lower()

        # Sliding window approach
        skill_len = len(skill)
        window_size = skill_len + 10  # Allow some flexibility

        for i in range(0, len(text_lower) - window_size + 1, 50):  # Step by 50 chars
            window = text_lower[i : i + window_size]
            ratio = fuzz.partial_ratio(skill_lower, window)
            if ratio > 85:
                spans.append({"start": i, "end": min(i + skill_len, len(text))})
                if len(spans) >= 3:
                    break

        return spans

    def _extract_years_for_skill(self, text: str, skill: str) -> Optional[float]:
        """Extract years of experience for a specific skill."""
        # Look for patterns like "5 years of Python" or "Python (3 years)"
        patterns = [
            rf"(\d+\.?\d*)\s*\+?\s*years?\s+(?:of\s+)?(?:experience\s+)?(?:with\s+)?{re.escape(skill)}",
            rf"(\d+\.?\d*)\s*\+?\s*years?\s+(?:working\s+)?(?:with\s+)?{re.escape(skill)}",
            rf"{re.escape(skill)}\s*[\(\[]?\s*(\d+\.?\d*)\s*\+?\s*years?",
            rf"{re.escape(skill)}.*?(\d+\.?\d*)\s*\+?\s*years?",
        ]

        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                try:
                    return float(match.group(1))
                except (ValueError, AttributeError):
                    pass

        return None

    async def _extract_skills_embeddings(
        self, text: str, existing_skills: Dict[str, str]
    ) -> Dict[str, Dict[str, Any]]:
        """Stage 2: Extract skills using embedding similarity."""
        try:
            # Generate embedding for the resume text
            response = await self.openai_client.embeddings.create(
                model="text-embedding-3-small", input=text[:8000]  # Limit text length
            )
            text_embedding = response.data[0].embedding

            # Get embeddings for vocabulary skills not yet found
            remaining_skills = [
                skill
                for skill in self.skills_vocab.values()
                if skill["canonical"] not in existing_skills
            ]

            if not remaining_skills:
                return {}

            # Batch generate embeddings for remaining skills
            skill_texts = [s["canonical"] for s in remaining_skills[:50]]  # Limit to 50
            skill_response = await self.openai_client.embeddings.create(
                model="text-embedding-3-small", input=skill_texts
            )

            # Calculate similarities
            found_skills = {}
            text_vec = np.array(text_embedding)

            for i, skill_data in enumerate(remaining_skills[:50]):
                skill_vec = np.array(skill_response.data[i].embedding)
                similarity = np.dot(text_vec, skill_vec) / (
                    np.linalg.norm(text_vec) * np.linalg.norm(skill_vec)
                )

                if similarity > 0.75:  # Threshold for semantic similarity
                    canonical = skill_data["canonical"]
                    found_skills[canonical] = {
                        "canonical": canonical,
                        "spans": [],  # No exact spans for embedding matches
                        "confidence": float(similarity),
                        "years": self._extract_years_for_skill(text, canonical),
                    }

            return found_skills

        except Exception as e:
            logger.error(f"Embedding extraction failed: {e}")
            return {}

    async def _extract_skills_openai(
        self, text: str, existing_skills: Dict[str, str]
    ) -> Dict[str, Dict[str, Any]]:
        """Stage 3: Extract skills using OpenAI function calling with closed-world constraint."""
        try:
            # Prepare vocabulary list for the prompt
            vocab_list = list(
                set([skill["canonical"] for skill in self.skills_vocab.values()])
            )

            # Create the function definition
            functions = [
                {
                    "name": "extract_skills",
                    "description": "Extract technical skills from resume text",
                    "parameters": {
                        "type": "object",
                        "properties": {
                            "skills": {
                                "type": "array",
                                "items": {
                                    "type": "object",
                                    "properties": {
                                        "skill": {
                                            "type": "string",
                                            "enum": vocab_list,  # Closed-world constraint
                                            "description": "Skill from the predefined vocabulary",
                                        },
                                        "confidence": {
                                            "type": "number",
                                            "minimum": 0,
                                            "maximum": 1,
                                            "description": "Confidence score (0-1)",
                                        },
                                        "years_experience": {
                                            "type": "number",
                                            "description": "Years of experience if mentioned",
                                        },
                                    },
                                    "required": ["skill", "confidence"],
                                },
                            }
                        },
                        "required": ["skills"],
                    },
                }
            ]

            # Load prompt template
            prompt_template = self._load_prompt_template()

            prompt = prompt_template.format(
                text=text[:4000],  # Limit text length
                existing_skills=(
                    ", ".join(existing_skills.keys()) if existing_skills else "None"
                ),
                vocabulary_sample=", ".join(
                    vocab_list[:50]
                ),  # Show sample of vocabulary
            )

            response = await self.openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {
                        "role": "system",
                        "content": "You are a technical recruiter expert at identifying skills in resumes. Only extract skills that are explicitly mentioned or strongly implied in the text.",
                    },
                    {"role": "user", "content": prompt},
                ],
                functions=functions,
                function_call={"name": "extract_skills"},
                temperature=0.3,
                max_tokens=1000,
            )

            # Parse the function call response
            if response.choices[0].message.function_call:
                import json

                result = json.loads(response.choices[0].message.function_call.arguments)

                found_skills = {}
                for skill_item in result.get("skills", []):
                    skill_name = skill_item["skill"]
                    if skill_name and skill_name not in existing_skills:
                        found_skills[skill_name] = {
                            "canonical": skill_name,
                            "spans": self._find_spans(text, skill_name),
                            "confidence": skill_item.get("confidence", 0.8),
                            "years": skill_item.get("years_experience"),
                        }

                return found_skills

            return {}

        except Exception as e:
            logger.error(f"OpenAI extraction failed: {e}")
            return {}

    def _load_prompt_template(self) -> str:
        """Load the skill extraction prompt template."""
        import os

        prompt_path = os.path.join(
            settings.CONFIG_DIR, "prompts", "skill_extraction.txt"
        )

        if os.path.exists(prompt_path):
            with open(prompt_path, "r") as f:
                return f.read()

        # Default prompt if file doesn't exist
        return """Extract technical skills from the following resume text. 
        
Resume Text:
{text}

Already Found Skills:
{existing_skills}

Available Vocabulary (sample):
{vocabulary_sample}

Instructions:
- Only extract skills that are explicitly mentioned or strongly implied
- Use ONLY skills from the predefined vocabulary
- Provide confidence scores based on how clearly the skill is mentioned
- Extract years of experience if mentioned near the skill
- Do not hallucinate or add skills not in the text
"""

    async def generate_embedding(self, text: str) -> List[float]:
        """Generate embedding vector for text using OpenAI."""
        try:
            response = await self.openai_client.embeddings.create(
                model="text-embedding-3-large",
                input=text[:8000],  # Limit to 8000 chars for token limits
                dimensions=3072,  # Match pgvector dimension
            )
            return response.data[0].embedding
        except Exception as e:
            logger.error(f"Failed to generate embedding: {e}")
            # Return zero vector as fallback
            return [0.0] * 3072
